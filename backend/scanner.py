"""HAVEN Universal Document Scanner — OCR + AI classification + auto-file.

Pipeline: upload (image/PDF) -> render to image -> Gemini 1.5 Flash (OCR + classify)
-> Apex Vault encrypt extracted PII -> auto-file into the Document Locker.

Primary AI remains local Ollama where applicable. Gemini free tier is used for
vision/document workloads when a local vision model is not configured.
"""
from __future__ import annotations
import base64
import json
import logging
import os
import uuid

from gemini_client import gemini_chat

logger = logging.getLogger("haven.scanner")

GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "")
SCANNER_MODEL_DISPLAY = "gemini-1.5-flash-latest"

CATEGORIES = [
    "identity",          # driver's license, state ID, passport
    "social_security",   # SSN card, SSA letters
    "birth_certificate",
    "medical",           # records, prescriptions, insurance cards
    "benefits",          # EBT/CalFresh, SSI/SSDI, Medi-Cal award letters
    "housing",           # lease, voucher, shelter letters, utility bills
    "income",            # pay stubs, W-2, employer letters
    "legal",             # court docs, citations, restraining orders
    "other",
]

SYSTEM_PROMPT = """You are BB, Haven's document intake assistant. You read scanned civic documents \
for residents experiencing housing/food/health insecurity. You perform OCR and classification. \
Respond ONLY with a single valid JSON object, no markdown fences, no prose."""

USER_PROMPT = f"""Read this document image carefully. Perform full OCR, then classify and extract.

Return ONLY this JSON object:
{{
  "category": "<one of: {', '.join(CATEGORIES)}>",
  "confidence": <0.0-1.0>,
  "title": "<short human title, e.g. 'California Driver License — John Smith'>",
  "summary": "<1-2 sentence plain-language summary of what this document is and why it matters>",
  "extracted_text": "<full OCR text of the document>",
  "key_fields": {{ "<field name>": "<value>" }},
  "expiration_date": "<YYYY-MM-DD or null>",
  "is_legible": <true|false>
}}

key_fields should capture identifiers like full name, document number, dates, issuing agency, case numbers, amounts.
If the image is not a document (random photo), use category "other", confidence 0, is_legible false."""


TEXT_PROMPT_TEMPLATE = """Read this document text carefully, then classify and extract. The full text is provided below between <document> tags.

Return ONLY the same JSON object schema as specified, where "extracted_text" is the cleaned document text.

<document>
{text}
</document>"""


def extract_docx_text(content: bytes) -> str:
    import io
    from docx import Document as DocxDocument
    d = DocxDocument(io.BytesIO(content))
    parts = [p.text for p in d.paragraphs if p.text.strip()]
    for table in d.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text.strip() for c in row.cells))
    return "\n".join(parts)


def extract_doc_text(content: bytes) -> str:
    import subprocess, tempfile
    with tempfile.NamedTemporaryFile(suffix=".doc", delete=False) as f:
        f.write(content)
        path = f.name
    try:
        out = subprocess.run(["antiword", path], capture_output=True, timeout=30)
        if out.returncode == 0 and out.stdout.strip():
            return out.stdout.decode("utf-8", errors="replace")
        raise ValueError("Could not read legacy .doc file — try saving it as .docx")
    finally:
        os.unlink(path)


def pdf_to_image_b64(pdf_bytes: bytes) -> str:
    """Render first page of a PDF to PNG base64 using PyMuPDF."""
    import fitz
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    page = doc.load_page(0)
    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
    png = pix.tobytes("png")
    doc.close()
    return base64.b64encode(png).decode("ascii")


async def scan_document(content: bytes, content_type: str, filename: str = "") -> dict:
    """Send document (image/PDF via vision, or txt/docx/doc as text) to Gemini; return classification dict."""
    if not GEMINI_API_KEY:
        raise RuntimeError("GEMINI_API_KEY not configured")

    ext = (filename or "").lower().rsplit(".", 1)[-1] if "." in (filename or "") else ""
    text_content = None
    if content_type == "text/plain" or ext == "txt":
        text_content = content.decode("utf-8", errors="replace")
    elif ext == "docx" or content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        text_content = extract_docx_text(content)
    elif ext == "doc" or content_type == "application/msword":
        text_content = extract_doc_text(content)

    chat = lambda system, user, image=None: gemini_chat(system, user, json_mode=True, image_base64=image)

    if text_content is not None:
        if not text_content.strip():
            raise ValueError("This file appears to be empty")
        raw = await chat(SYSTEM_PROMPT, USER_PROMPT + "\n\n" + TEXT_PROMPT_TEMPLATE.format(text=text_content[:30000]))
    else:
        if content_type == "application/pdf":
            image_b64 = pdf_to_image_b64(content)
        else:
            image_b64 = base64.b64encode(content).decode("ascii")
        raw = await chat(SYSTEM_PROMPT, USER_PROMPT, image_base64=image_b64)
    text = str(raw).strip()
    if text.startswith("```"):
        text = text.strip("`")
        if text.startswith("json"):
            text = text[4:]
    try:
        result = json.loads(text)
    except json.JSONDecodeError:
        start, end = text.find("{"), text.rfind("}")
        if start == -1 or end == -1:
            raise ValueError(f"Scanner returned non-JSON output: {text[:200]}")
        result = json.loads(text[start:end + 1])

    if result.get("category") not in CATEGORIES:
        result["category"] = "other"
    result["confidence"] = float(result.get("confidence") or 0)
    return result
